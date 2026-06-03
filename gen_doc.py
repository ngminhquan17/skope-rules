"""Script tạo tài liệu DOCX mô tả toàn bộ ý tưởng và logic của SkopeRules."""

import sys
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants


def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Times New Roman"
    return p


def body(doc, text, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=12)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 1.0)
    run = p.add_run(text)
    set_font(run, size=12)
    return p


def code_block(doc, lines):
    """Thêm block code dạng monospace."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x6E)


def section_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=(0x1A, 0x53, 0x76))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Shading Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            for run in cells[c_idx].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
    return table


# ============================================================
# TẠO TÀI LIỆU
# ============================================================

doc = Document()

# --- Margin ---
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)

# ============================================================
# TRANG TIÊU ĐỀ
# ============================================================

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SKOPE RULES")
set_font(run, size=24, bold=True, color=(0x1A, 0x53, 0x76))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Tài liệu Ý tưởng & Logic Thực hiện")
set_font(run, size=16, italic=True, color=(0x44, 0x44, 0x44))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Phiên bản đang phát triển  ·  2026")
set_font(run, size=11, color=(0x88, 0x88, 0x88))

doc.add_page_break()

# ============================================================
# MỤC LỤC (thủ công)
# ============================================================

heading(doc, "Mục lục", level=1)
toc_items = [
    ("1.", "Tổng quan dự án"),
    ("2.", "Động lực & Bài toán cần giải quyết"),
    ("3.", "Kiến trúc tổng thể"),
    ("4.", "Bước 1 – Huấn luyện ensemble cây quyết định"),
    ("5.", "Bước 2 – Trích xuất rule từ cây (Tree-to-Rules)"),
    ("6.", "Bước 3 – Đánh giá hiệu suất rule trên OOB"),
    ("7.", "Bước 4 – Lọc, factorize và deduplicate rule"),
    ("8.", "Bước 5 – Tính Activation Set cho mỗi rule"),
    ("9.", "Bước 6 – Đo độ tương đồng giữa các rule"),
    ("10.", "Bước 7 – Phân cụm rule theo đồ thị (Graph Clustering)"),
    ("11.", "Bước 8 – Ensemble Scoring (đang phát triển)"),
    ("12.", "Lớp Rule & Logic Factorize"),
    ("13.", "Tham số cấu hình"),
    ("14.", "Đầu ra của mô hình"),
    ("15.", "Hướng phát triển tiếp theo"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"{num}  {title}")
    set_font(run, size=12)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. TỔNG QUAN
# ============================================================

heading(doc, "1. Tổng quan dự án", level=1)

body(doc,
    "SkopeRules là một thư viện Python dành cho bài toán phát hiện bất thường "
    "(anomaly detection) và phân loại có khả năng diễn giải (interpretable "
    "classification). Mục tiêu cốt lõi là tự động sinh ra các rule logic đơn giản, "
    "dễ đọc, ở dạng mà con người có thể hiểu và kiểm chứng được — chẳng hạn:"
)

doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.5)
run = p.add_run(
    '"worst radius > 16.8  AND  worst concavity > 0.21"  →  khả năng cao là khối u ác tính'
)
set_font(run, size=12, italic=True, color=(0x1A, 0x53, 0x76))

body(doc,
    "Thư viện tuân thủ giao diện sklearn (BaseEstimator), có thể dùng ngay trong "
    "pipeline machine learning thông thường bằng fit() / predict()."
)

heading(doc, "Công nghệ sử dụng", level=2)
add_table(doc,
    ["Thành phần", "Thư viện / Kỹ thuật"],
    [
        ("Mô hình nền",          "scikit-learn BaggingClassifier + DecisionTreeClassifier"),
        ("Xử lý dữ liệu",        "NumPy, Pandas"),
        ("Phân cụm rule",        "Đồ thị vô hướng + DFS (tự cài đặt)"),
        ("Đo tương đồng rule",   "Jaccard / Adjusted Jaccard / Asymmetric Similarity"),
        ("Giao diện mô hình",    "sklearn BaseEstimator"),
    ]
)

doc.add_paragraph()

# ============================================================
# 2. ĐỘNG LỰC
# ============================================================

heading(doc, "2. Động lực & Bài toán cần giải quyết", level=1)

body(doc,
    "Các mô hình machine learning hiện đại (Random Forest, Gradient Boosting, Neural "
    "Network) đạt hiệu suất cao nhưng hoàn toàn là hộp đen — không ai biết tại sao "
    "một mẫu bị gán nhãn 'bất thường'. Trong nhiều lĩnh vực thực tế (y tế, tài chính, "
    "phát hiện gian lận), đây là vấn đề nghiêm trọng:"
)

for item in [
    "Bác sĩ cần biết lý do trước khi tin vào chẩn đoán của AI.",
    "Ngân hàng phải giải thích tại sao từ chối một giao dịch.",
    "Kỹ sư vận hành cần hiểu điều kiện nào khiến thiết bị hỏng.",
]:
    bullet(doc, item)

body(doc,
    "SkopeRules giải quyết điều này bằng cách chuyển đổi từ cây quyết định sang "
    "rule ngôn ngữ tự nhiên, đồng thời chọn lọc những rule thực sự có ý nghĩa thống kê "
    "(precision ≥ ngưỡng, recall ≥ ngưỡng) và nhóm các rule tương đồng lại với nhau "
    "để dễ phân tích."
)

# ============================================================
# 3. KIẾN TRÚC TỔNG THỂ
# ============================================================

heading(doc, "3. Kiến trúc tổng thể", level=1)

body(doc,
    "Toàn bộ pipeline của SkopeRules gồm 2 file chính và 8 bước tuần tự trong fit():"
)

add_table(doc,
    ["File", "Vai trò"],
    [
        ("skope_rules.py", "Lớp SkopeRules — toàn bộ pipeline huấn luyện và dự đoán"),
        ("rule.py",        "Lớp Rule — phân tích cú pháp, factorize, tính activation mask"),
    ]
)

doc.add_paragraph()
body(doc, "Sơ đồ pipeline (từ trái sang phải):")

steps = [
    "Dữ liệu X, y",
    "Huấn luyện BaggingClassifier (nhiều cây)",
    "Trích xuất rule từ từng cây (DFS)",
    "Đánh giá precision / recall trên OOB",
    "Lọc rule (precision_min, recall_min)",
    "Factorize + Deduplicate (Rule class)",
    "Tính Activation Set S_r cho mỗi rule",
    "Đo tương đồng (Adjusted Jaccard)",
    "Graph Clustering (DFS components)",
    "Kết quả: clustered_rules_",
]
for i, s in enumerate(steps):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    arrow = "→  " if i > 0 else "   "
    run = p.add_run(f"{arrow}[{i}] {s}")
    set_font(run, size=11, bold=(i == 0 or i == len(steps)-1))
    p.paragraph_format.space_after = Pt(1)

doc.add_paragraph()

# ============================================================
# 4. BƯỚC 1 — ENSEMBLE CÂY
# ============================================================

heading(doc, "4. Bước 1 – Huấn luyện ensemble cây quyết định", level=1)

body(doc,
    "Mỗi BaggingClassifier bao gồm n_estimators cây quyết định "
    "DecisionTreeClassifier. Dữ liệu huấn luyện được lấy mẫu ngẫu nhiên "
    "(bootstrap=True để hỗ trợ OOB evaluation). Quá trình này tạo ra sự đa dạng "
    "giữa các cây, từ đó sinh ra nhiều loại rule khác nhau."
)

heading(doc, "Tại sao dùng BaggingClassifier?", level=2)
for item in [
    "Mỗi cây chỉ thấy một phần dữ liệu → các path dẫn đến leaf node sẽ khác nhau → nhiều rule đa dạng.",
    "Bootstrap sampling tạo ra tập Out-Of-Bag (OOB) — dùng để đánh giá khách quan mà không cần tập test riêng.",
    "Nhiều cây với max_depth khác nhau (truyền vào dạng list) cho phép tạo rule ngắn và dài song song.",
]:
    bullet(doc, item)

heading(doc, "Tham số quan trọng", level=2)
add_table(doc,
    ["Tham số", "Ý nghĩa"],
    [
        ("n_estimators",        "Số cây trong mỗi BaggingClassifier"),
        ("max_samples",         "Tỷ lệ mẫu lấy ngẫu nhiên cho mỗi cây (mặc định 80%)"),
        ("max_depth",           "Độ sâu tối đa của cây — kiểm soát độ phức tạp của rule"),
        ("max_features",        "Số feature mỗi lần split bên trong cây"),
        ("max_samples_features","Tỷ lệ feature lấy ngẫu nhiên cho mỗi cây trong Bagging"),
        ("bootstrap",           "True để bật OOB evaluation (khuyến nghị)"),
    ]
)

doc.add_paragraph()
body(doc, "Code khởi tạo (đã sửa so với code gốc có lỗi):")
code_block(doc, [
    "bagging_clf = BaggingClassifier(",
    "    estimator=DecisionTreeClassifier(",
    "        max_depth=max_depth,",
    "        max_features=self.max_features,",
    "        min_samples_split=self.min_samples_split,",
    "    ),",
    "    n_estimators=self.n_estimators,",
    "    max_samples=max_samples,",
    "    bootstrap=self.bootstrap,",
    "    random_state=self.random_state,",
    ")",
])

# ============================================================
# 5. BƯỚC 2 — TREE TO RULES
# ============================================================

heading(doc, "5. Bước 2 – Trích xuất rule từ cây (Tree-to-Rules)", level=1)

body(doc,
    "Sau khi mỗi cây được huấn luyện, toàn bộ path từ gốc đến mỗi leaf node "
    "được chuyển thành một rule điều kiện. Hàm _tree_to_rules() duyệt cây theo "
    "DFS đệ quy:"
)

bullet(doc, "Khi gặp node trong (internal node): thêm điều kiện 'feature <= threshold' vào nhánh trái, 'feature > threshold' vào nhánh phải.")
bullet(doc, "Khi gặp leaf node: ghép toàn bộ điều kiện trên path bằng ' and ' → tạo thành một rule hoàn chỉnh.")
bullet(doc, "Một cây có max_depth = d sẽ tạo ra tối đa 2^d rule.")

body(doc, "Ví dụ với cây có độ sâu 2:")
code_block(doc, [
    "Cây:",
    "  [worst_radius <= 16.8]",
    "        /          \\",
    " [area <= 500]   → Rule B: 'worst_radius > 16.8'",
    "    /      \\",
    "leaf      leaf",
    "",
    "Rule A: 'worst_radius <= 16.8 AND area <= 500'",
    "Rule C: 'worst_radius <= 16.8 AND area > 500'",
    "Rule B: 'worst_radius > 16.8'",
])

body(doc,
    "Lưu ý: feature name trong rule ở bước này vẫn dùng tên nội bộ "
    "(__C__0, __C__1, ...) để tránh xung đột ký tự khi dùng pandas.query(). "
    "Tên thật được thay thế ở bước cuối bằng replace_feature_name()."
)

# ============================================================
# 6. BƯỚC 3 — OOB EVALUATION
# ============================================================

heading(doc, "6. Bước 3 – Đánh giá hiệu suất rule trên OOB", level=1)

body(doc,
    "Với bootstrap=True, mỗi cây chỉ được huấn luyện trên ~63.2% dữ liệu. "
    "Phần còn lại (~36.8%) là Out-Of-Bag (OOB) — chưa bao giờ được cây này thấy. "
    "Đây là tập validation tự nhiên, không cần tách riêng."
)

heading(doc, "Hai chỉ số được tính", level=2)
add_table(doc,
    ["Chỉ số", "Công thức", "Ý nghĩa"],
    [
        ("Precision",
         "TP / (TP + FP)",
         "Trong các mẫu rule bắt được, bao nhiêu % thực sự là anomaly?"),
        ("Recall",
         "TP / (TP + FN)",
         "Rule bắt được bao nhiêu % tổng số anomaly thực?"),
    ]
)

doc.add_paragraph()
body(doc, "Hàm _eval_rule_perf() thực hiện:")
code_block(doc, [
    "detected_index = X_oob.query(rule).index    # pandas lọc mẫu thỏa rule",
    "y_detected     = y_oob[detected_index]",
    "true_pos       = y_detected[y_detected > 0].sum()",
    "precision      = y_detected.mean()           # tỷ lệ dương tính đúng",
    "recall         = true_pos / total_positives  # tỷ lệ bao phủ anomaly",
])

body(doc,
    "Khi một rule xuất hiện nhiều lần (từ nhiều cây khác nhau), precision và recall "
    "của nó được cập nhật theo trung bình chạy (running mean) — tránh overfit vào "
    "một cây đơn lẻ."
)

# ============================================================
# 7. BƯỚC 4 — LỌC VÀ FACTORIZE
# ============================================================

heading(doc, "7. Bước 4 – Lọc, Factorize và Deduplicate rule", level=1)

heading(doc, "4a. Lọc theo ngưỡng", level=2)
body(doc,
    "Chỉ giữ lại các rule thỏa mãn đồng thời hai điều kiện:"
)
bullet(doc, "precision >= precision_min  (mặc định 0.5)")
bullet(doc, "recall    >= recall_min     (mặc định 0.01)")

heading(doc, "4b. Factorize — loại bỏ điều kiện thừa", level=2)
body(doc,
    "Lớp Rule phân tích cú pháp rule string và đơn giản hóa các điều kiện trùng "
    "lặp trên cùng một feature:"
)
code_block(doc, [
    "# Trước factorize:",
    "'worst_radius > 10.0 AND worst_radius > 16.5 AND area <= 500'",
    "",
    "# Sau factorize (giữ lại điều kiện chặt hơn):",
    "'area <= 500.0 AND worst_radius > 16.5'",
])
body(doc,
    "Kết quả: rule ngắn gọn hơn, đồng thời dùng được làm key trong dict để deduplicate "
    "(vì __hash__ và __eq__ của Rule class dựa trên agg_dict đã factorize)."
)

heading(doc, "4c. Sắp xếp rule theo chất lượng", level=2)
body(doc,
    "Danh sách rule cuối được sắp xếp giảm dần theo (precision, recall) — "
    "rule tốt nhất đứng đầu."
)

# ============================================================
# 8. BƯỚC 5 — ACTIVATION SET
# ============================================================

heading(doc, "8. Bước 5 – Tính Activation Set cho mỗi rule", level=1)

body(doc,
    "Với mỗi rule r, ta tính một vector nhị phân S_r ∈ {0,1}^N, "
    "trong đó S_r[i] = 1 nếu mẫu thứ i thỏa mãn tất cả điều kiện của rule r."
)

code_block(doc, [
    "def evaluate(self, X, feature_names):",
    "    mask = np.ones(len(X), dtype=bool)    # ban đầu chọn tất cả",
    "    for feature, symbol in self.agg_dict:",
    "        value = float(self.agg_dict[(feature, symbol)])",
    "        idx   = feature_names.index(feature)",
    "        col   = X[:, idx]",
    "        if symbol == '>'  : mask &= (col > value)",
    "        if symbol == '<=' : mask &= (col <= value)",
    "        # ...các trường hợp còn lại",
    "    return mask",
])

body(doc,
    "Activation Set chính là cơ sở để so sánh hai rule: nếu hai rule bắt gần như "
    "cùng một tập mẫu → chúng nên nằm trong cùng cluster."
)

# ============================================================
# 9. BƯỚC 6 — ĐO TƯƠNG ĐỒNG
# ============================================================

heading(doc, "9. Bước 6 – Đo độ tương đồng giữa các rule", level=1)

body(doc,
    "Ba hàm đo tương đồng được cài đặt, mỗi hàm phù hợp với một góc nhìn khác nhau:"
)

heading(doc, "9a. Jaccard Similarity (cơ bản)", level=2)
code_block(doc, [
    "J(r1, r2) = |S_r1 ∩ S_r2| / |S_r1 ∪ S_r2|",
])
body(doc,
    "Nhanh, đơn giản. Nhược điểm: rule rộng (bắt nhiều mẫu) luôn có Jaccard cao "
    "với mọi rule khác — dễ gây ra cluster giả."
)

heading(doc, "9b. Adjusted Jaccard (đang dùng)", level=2)
code_block(doc, [
    "adj_sim(r1, r2) = J(r1, r2) × (1 − max(coverage(r1), coverage(r2)))",
    "",
    "coverage(r) = |S_r| / N    # tỷ lệ mẫu mà rule bắt được",
])
body(doc,
    "Nhân thêm hệ số phạt khi một trong hai rule bắt quá nhiều mẫu (coverage cao). "
    "Giúp tránh việc một rule rộng kéo tất cả rule khác vào cùng cluster."
)

heading(doc, "9c. Asymmetric Similarity (phát hiện quan hệ con)", level=2)
code_block(doc, [
    "asym(r1, r2) = |S_r1 ∩ S_r2| / min(|S_r1|, |S_r2|)",
])
body(doc,
    "Phát hiện khi rule nhỏ nằm hoàn toàn trong rule lớn hơn "
    "(ví dụ: r1 là trường hợp đặc biệt của r2). "
    "Hiện được cài đặt nhưng chưa dùng trong pipeline chính."
)

# ============================================================
# 10. BƯỚC 7 — CLUSTERING
# ============================================================

heading(doc, "10. Bước 7 – Phân cụm rule theo đồ thị (Graph Clustering)", level=1)

body(doc,
    "Sau khi tính ma trận tương đồng, thuật toán xây dựng đồ thị vô hướng G = (V, E):"
)
bullet(doc, "Mỗi vertex V = một rule.")
bullet(doc, "Mỗi cạnh E giữa hai rule được thêm nếu adjusted_similarity > threshold (mặc định 0.6).")

body(doc, "Sau đó tìm các connected component bằng DFS:")
code_block(doc, [
    "visited = set()",
    "for i in range(n_rules):",
    "    if i in visited: continue",
    "    stack = [i]",
    "    comp  = []",
    "    while stack:",
    "        node = stack.pop()",
    "        if node in visited: continue",
    "        visited.add(node)",
    "        comp.append(node)",
    "        stack.extend(adj[node])   # thêm hàng xóm vào stack",
    "    clusters.append(comp)",
])

body(doc,
    "Mỗi connected component là một cluster. Rule trong cùng cluster có "
    "pattern tương đồng — bắt cùng nhóm mẫu bất thường — nhưng có thể "
    "dùng feature hoặc ngưỡng khác nhau."
)

heading(doc, "Ý nghĩa của clustering", level=2)
for item in [
    "Nhóm các rule 'nói cùng một điều' để người dùng có thể chọn rule đại diện.",
    "Mỗi cluster ứng với một 'pattern bất thường' khác nhau trong dữ liệu.",
    "Giảm số lượng rule cần đọc: thay vì đọc 50 rule, người dùng chỉ cần đọc 5-10 cluster.",
    "max_cluster_size giới hạn kích thước cluster để tránh cluster khổng lồ.",
]:
    bullet(doc, item)

# ============================================================
# 11. BƯỚC 8 — ENSEMBLE SCORING
# ============================================================

heading(doc, "11. Bước 8 – Ensemble Scoring (đang phát triển)", level=1)

body(doc,
    "Thay vì dùng rule tốt nhất để dự đoán (chọn max), _ensemble_score() "
    "tổng hợp tín hiệu từ tất cả rule trong một cluster bằng cách vote có trọng số:"
)
code_block(doc, [
    "score = Σ (weight_r × S_r)  /  Σ weight_r",
    "",
    "# Mặc định: weight_r = precision của rule r",
    "# score[i] ≈ xác suất trung bình mẫu i là anomaly theo cluster này",
])
body(doc,
    "Điểm số này mịn hơn hard-vote và ít nhạy cảm với nhiễu hơn so với việc "
    "chọn một rule duy nhất. Phương pháp này chưa được tích hợp vào predict() "
    "trong phiên bản hiện tại."
)

# ============================================================
# 12. LỚP RULE
# ============================================================

heading(doc, "12. Lớp Rule & Logic Factorize", level=1)

body(doc,
    "Rule class đóng vai trò quan trọng: chuẩn hóa, so sánh và đánh giá rule."
)

add_table(doc,
    ["Method", "Chức năng"],
    [
        ("__init__(rule, args)",  "Phân tích rule string thành dict (feature, symbol) → value"),
        ("factorize()",           "Loại bỏ điều kiện thừa: X > 3 AND X > 5 → X > 5"),
        ("__eq__ / __hash__",     "Dùng agg_dict đã factorize → hai rule giống nhau sau simplify sẽ bằng nhau"),
        ("__repr__",              "In rule dạng chuẩn (sắp xếp theo tên feature)"),
        ("evaluate(X, names)",    "Trả về boolean mask — mẫu nào thỏa rule"),
        ("__iter__",              "Cho phép unpack: rule_str, args = Rule(...)"),
    ]
)

doc.add_paragraph()
body(doc, "Logic factorize cho từng loại điều kiện:")
add_table(doc,
    ["Điều kiện", "Giữ lại"],
    [
        ("X > 3  AND  X > 7",   "X > 7     (điều kiện chặt hơn phía lớn hơn)"),
        ("X <= 10  AND  X <= 6", "X <= 6    (điều kiện chặt hơn phía nhỏ hơn)"),
        ("X > 3  AND  X <= 10",  "Giữ cả hai (giới hạn hai phía)"),
    ]
)

# ============================================================
# 13. THAM SỐ CẤU HÌNH
# ============================================================

heading(doc, "13. Tham số cấu hình", level=1)

add_table(doc,
    ["Tham số", "Mặc định", "Ảnh hưởng"],
    [
        ("precision_min",        "0.5",   "Ngưỡng lọc rule — tăng lên để chỉ giữ rule chắc chắn hơn"),
        ("recall_min",           "0.01",  "Ngưỡng coverage — giảm để giữ rule bắt ít mẫu nhưng chính xác"),
        ("n_estimators",         "10",    "Số cây — nhiều hơn → nhiều rule hơn, chậm hơn"),
        ("max_depth",            "6",     "Độ sâu cây — nông → rule ngắn, dễ hiểu; sâu → rule chi tiết hơn"),
        ("max_depth=[2,3,4]",    "—",     "Train nhiều BaggingClassifier với độ sâu khác nhau đồng thời"),
        ("max_samples",          "0.8",   "Tỷ lệ mẫu mỗi cây — phải < 1 để có OOB"),
        ("bootstrap",            "False", "Bật True để dùng OOB evaluation chuẩn"),
        ("threshold (cluster)",  "0.6",   "Ngưỡng adjusted_similarity để nối hai rule trong đồ thị"),
    ]
)

doc.add_paragraph()

# ============================================================
# 14. ĐẦU RA
# ============================================================

heading(doc, "14. Đầu ra của mô hình", level=1)

body(doc,
    "Sau khi gọi fit(), model lưu các thuộc tính sau:"
)

add_table(doc,
    ["Thuộc tính", "Kiểu dữ liệu", "Mô tả"],
    [
        ("rules_",
         "list of (str, tuple)",
         "Danh sách (rule_string, (precision, recall, count)) sắp xếp theo precision giảm dần"),
        ("clusters_",
         "list of list of int",
         "Chỉ số của rule trong rules_ theo từng cluster"),
        ("clustered_rules_",
         "list of list of str",
         "Rule string dạng đọc được (đã thay tên feature thật) theo từng cluster"),
        ("clustered_rules_raw_",
         "list of list of str",
         "Rule string dùng tên nội bộ __C__0, __C__1 (dùng cho pandas query)"),
        ("estimators_",
         "list",
         "Toàn bộ cây quyết định đã huấn luyện"),
        ("feature_dict_",
         "dict",
         "Bảng ánh xạ __C__0 → tên feature thật"),
    ]
)

doc.add_paragraph()
body(doc, "Ví dụ đầu ra clustered_rules_ trên tập Breast Cancer:")
code_block(doc, [
    "Cluster 1  [10 rules]  — pattern: 'worst radius lớn + worst concavity cao'",
    "  (1) worst radius > 16.8 AND worst concavity > 0.21",
    "       precision=1.00  recall=0.80  support≈118 mẫu",
    "  (2) worst perimeter > 112.8 AND worst area > 810 AND worst concavity > 0.21",
    "       precision=0.98  recall=0.83  support≈123 mẫu",
    "  ...",
    "",
    "Cluster 2  [1 rule]  — pattern khác biệt",
    "  (1) mean texture > 15.9 AND worst concave points > 0.15",
    "       precision=1.00  recall=0.63  support≈93 mẫu",
])

# ============================================================
# 15. HƯỚNG PHÁT TRIỂN
# ============================================================

heading(doc, "15. Hướng phát triển tiếp theo", level=1)

items = [
    ("Tích hợp Ensemble Scoring vào predict()",
     "Dùng _ensemble_score() theo từng cluster thay vì chọn rule đơn lẻ. "
     "Kết quả dự đoán sẽ mịn hơn và ổn định hơn."),
    ("Chọn rule đại diện cho mỗi cluster",
     "Tự động chọn một rule duy nhất từ mỗi cluster (ví dụ: rule có F1 cao nhất) "
     "để rút gọn output từ N rule → K cluster representative."),
    ("Asymmetric Similarity vào pipeline",
     "Dùng _asymmetric_similarity() để phát hiện quan hệ cha-con giữa các rule, "
     "từ đó xây dựng hierarchy rule thay vì flat clustering."),
    ("Hỗ trợ multi-class",
     "Mở rộng từ binary (0/1) sang phát hiện nhiều loại anomaly khác nhau."),
    ("Đóng gói thành Python package",
     "Thêm __init__.py, setup.py / pyproject.toml, tests và đưa lên PyPI."),
    ("Visualization",
     "Vẽ đồ thị cluster rule, highlight feature quan trọng nhất trong mỗi cluster."),
]

for title, desc in items:
    section_title(doc, f"▸  {title}")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(desc)
    set_font(run, size=12)
    p.paragraph_format.space_after = Pt(6)

# ============================================================
# LƯU FILE
# ============================================================

out_path = r"d:\FUTURE\skope_rules\SkopeRules_Project_Document.docx"
doc.save(out_path)
print(f"Đã lưu tài liệu: {out_path}")
